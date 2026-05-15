from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Dict, List

import pdfplumber
from docx import Document

from services.client_registry import find_known_client
from services.domain_detector import detect_domain
from utils.schema import Experience, ResumeProfile
from utils.text import clean_text, dedupe_keep_order, split_csvish


SECTION_PATTERNS = {
    "professional_summary": r"(professional summary|summary|profile)",
    "technical_skills": r"(technical skills|skills|technologies)",
    "experience": r"(professional experience|work experience|experience|client experience)",
    "certifications": r"(certifications?|licenses?)",
    "education": r"(education|academic)",
}

def extract_resume_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError("Unsupported resume type. Upload a DOCX or PDF file.")


def _extract_pdf(data: bytes) -> str:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = [page.extract_text(x_tolerance=1, y_tolerance=3) or "" for page in pdf.pages]
    return clean_text("\n".join(pages))


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    lines: List[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return clean_text("\n".join(lines))


def parse_resume_text(text: str) -> ResumeProfile:
    text = clean_text(text)
    sections = _split_sections(text)
    profile = ResumeProfile(raw_text=text, sections=sections)
    profile.professional_summary = sections.get("professional_summary", "")
    profile.technical_skills = _parse_skill_section(sections.get("technical_skills", ""))
    profile.experiences = _parse_experiences(sections.get("experience", text), text)
    profile.certifications = _parse_simple_lines(sections.get("certifications", ""))
    profile.education = sections.get("education", "")
    return profile


def load_prompt(name: str) -> str:
    prompt_path = Path(__file__).resolve().parents[1] / "prompts" / name
    return prompt_path.read_text(encoding="utf-8")


def _split_sections(text: str) -> Dict[str, str]:
    lines = text.splitlines()
    section_by_line: List[tuple[int, str]] = []
    for idx, line in enumerate(lines):
        normalized = re.sub(r"[^a-zA-Z ]", "", line).strip().lower()
        for section, pattern in SECTION_PATTERNS.items():
            if re.fullmatch(pattern, normalized, flags=re.IGNORECASE):
                section_by_line.append((idx, section))
                break

    sections: Dict[str, str] = {}
    for pos, (line_idx, section) in enumerate(section_by_line):
        next_idx = section_by_line[pos + 1][0] if pos + 1 < len(section_by_line) else len(lines)
        content = "\n".join(lines[line_idx + 1 : next_idx]).strip()
        sections[section] = content
    return sections


def _parse_skill_section(text: str) -> Dict[str, List[str]]:
    if not text:
        return {}
    skills: Dict[str, List[str]] = {}
    current = "Core Skills"
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if ":" in line and len(line.split(":", 1)[0].split()) <= 5:
            heading, values = line.split(":", 1)
            current = heading.strip()
            skills[current] = split_csvish(values)
        else:
            skills.setdefault(current, []).extend(split_csvish(line))
    return {k: dedupe_keep_order(v) for k, v in skills.items() if v}


def _parse_experiences(text: str, full_text: str) -> List[Experience]:
    chunks = _experience_chunks(text)
    experiences: List[Experience] = []
    for chunk in chunks:
        client = _find_client(chunk)
        title = _find_title(chunk)
        dates = _find_dates(chunk)
        responsibilities = _find_bullets(chunk)
        environment = _find_environment(chunk)
        domain = detect_domain(client, chunk or full_text)
        experiences.append(
            Experience(
                client_name=client or "Client",
                title=title,
                dates=dates,
                domain=domain,
                responsibilities=responsibilities,
                environment=environment,
                raw_text=chunk,
            )
        )
    return experiences


def _experience_chunks(text: str) -> List[str]:
    known_client_chunks = _known_client_chunks(text)
    if len(known_client_chunks) >= 2:
        return known_client_chunks

    client_matches = list(re.finditer(r"(?im)^\s*(?:client\s*[:|-]\s*)?([A-Za-z0-9&.,'() -]{3,})\s*$", text))
    filtered = [m for m in client_matches if _looks_like_client_line(m.group(0))]
    if len(filtered) >= 2:
        chunks: List[str] = []
        for idx, match in enumerate(filtered):
            start = match.start()
            end = filtered[idx + 1].start() if idx + 1 < len(filtered) else len(text)
            chunks.append(text[start:end].strip())
        return chunks

    date_matches = list(re.finditer(r"(?i)(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*[-to]+\s*(present|\d{4}|[a-z]+\.?\s+\d{4})", text))
    if len(date_matches) >= 2:
        chunks = []
        for idx, match in enumerate(date_matches):
            start = max(0, text.rfind("\n\n", 0, match.start()))
            end = date_matches[idx + 1].start() if idx + 1 < len(date_matches) else len(text)
            chunks.append(text[start:end].strip())
        return chunks
    return [text] if text else []


def _known_client_chunks(text: str) -> List[str]:
    matches: List[int] = []
    seen_clients = set()
    offset = 0

    for line in text.splitlines(keepends=True):
        stripped = line.strip()
        client = find_known_client(stripped)
        if client and _is_probable_client_heading(stripped) and client.name.lower() not in seen_clients:
            matches.append(offset)
            seen_clients.add(client.name.lower())
        offset += len(line)

    if len(matches) < 2:
        return []

    chunks: List[str] = []
    for idx, start in enumerate(matches):
        end = matches[idx + 1] if idx + 1 < len(matches) else len(text)
        chunks.append(text[start:end].strip())
    return chunks


def _is_probable_client_heading(line: str) -> bool:
    lowered = line.lower()
    if not line or line.startswith(("-", "*", "•")):
        return False
    if any(term in lowered for term in ["environment", "responsibilities", "summary", "skills", "education"]):
        return False
    return len(line) <= 120


def _looks_like_client_line(line: str) -> bool:
    lower = line.lower()
    if any(word in lower for word in ["summary", "skills", "education", "certification", "environment"]):
        return False
    return find_known_client(line) is not None


def _find_client(chunk: str) -> str:
    explicit = re.search(r"(?im)^\s*client\s*[:|-]\s*(.+?)\s*$", chunk)
    if explicit:
        known = find_known_client(explicit.group(1))
        return known.name if known else explicit.group(1).strip()
    known = find_known_client(chunk)
    if known:
        return known.name
    for line in chunk.splitlines()[:4]:
        clean = line.strip(" |:-")
        if clean and not _find_dates(clean) and len(clean.split()) <= 8:
            return clean
    return "Client"


def _find_title(chunk: str) -> str:
    titles = [
        "Senior Data Engineer",
        "Data Engineer",
        "Big Data Engineer",
        "ETL Developer",
        "Software Engineer",
        "Data Analyst",
    ]
    for title in titles:
        if re.search(title, chunk, flags=re.IGNORECASE):
            return title
    return "Data Engineer"


def _find_dates(chunk: str) -> str:
    patterns = [
        r"(?i)(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*[-to]+\s*(present|\d{4}|[a-z]+\.?\s+\d{4})",
        r"\b(20\d{2}|19\d{2})\s*[-to]+\s*(present|20\d{2}|19\d{2})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, chunk)
        if match:
            return match.group(0).strip()
    return ""


def _find_bullets(chunk: str) -> List[str]:
    bullets = []
    for line in chunk.splitlines():
        stripped = line.strip()
        if stripped.startswith(("-", "•", "*")):
            bullets.append(stripped.lstrip("-•* ").strip())
    return dedupe_keep_order(bullets)


def _find_environment(chunk: str) -> List[str]:
    env_match = re.search(r"(?is)environment\s*[:|-]\s*(.+?)(?:\n\n|$)", chunk)
    if not env_match:
        return []
    return split_csvish(env_match.group(1))


def _parse_simple_lines(text: str) -> List[str]:
    return dedupe_keep_order(line.strip("-•* ") for line in text.splitlines() if line.strip())
