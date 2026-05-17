from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from utils.docx_exporter import build_docx
from utils.schema import Experience, TailoredResume


class DocxExporterBoldFormattingTests(unittest.TestCase):
    def test_technical_skill_values_are_plain_while_category_labels_are_bold(self) -> None:
        document = _build_sample_document()

        skill_paragraphs = _section_paragraphs(document, "Technical Skills", "Professional Experience")
        self.assertGreaterEqual(len(skill_paragraphs), 2)

        for paragraph in skill_paragraphs:
            self.assertEqual(paragraph.style.name, "Normal")
            self.assertEqual(paragraph.paragraph_format.space_after, Pt(0))
            self.assertEqual(paragraph.paragraph_format.line_spacing, 1.0)
            self.assertTrue(paragraph.runs[0].bold)
            self.assertTrue(paragraph.runs[0].text.endswith(": "))

            value_runs = [run for run in paragraph.runs[1:] if run.text.strip()]
            self.assertTrue(value_runs)
            self.assertEqual(paragraph.runs[0].font.size, Pt(10))
            self.assertTrue(all(run.bold is not True for run in value_runs))
            self.assertTrue(all(run.font.size == Pt(10) for run in value_runs))

    def test_experience_bullets_only_bold_known_environment_technologies(self) -> None:
        document = _build_sample_document()

        bullet = _paragraph_containing(document, "patient analytics with 35% faster order pipelines")
        bold_text = _bold_text(bullet)

        self.assertIn("AWS", bold_text)
        self.assertIn("Glue", bold_text)
        self.assertIn("SQL", bold_text)
        self.assertNotIn("patient analytics", bold_text)
        self.assertNotIn("35%", bold_text)
        self.assertNotIn("order pipelines", bold_text)

    def test_environment_section_bolds_only_label_not_tools(self) -> None:
        document = _build_sample_document()

        environment = _paragraph_containing(document, "Environment: AWS, Glue")
        bold_runs = [run.text for run in environment.runs if run.bold is True]
        environment_label = next(run for run in environment.runs if run.text == "Environment: ")
        value_runs = [run for run in environment.runs if run.text.strip() and run.text != "Environment: "]
        bold_text = "".join(bold_runs)

        self.assertIn("Environment: ", bold_runs)
        self.assertEqual(environment_label.font.size, Pt(10))
        self.assertTrue(all(run.font.size == Pt(10) for run in value_runs))
        self.assertNotIn("AWS", bold_text)
        self.assertNotIn("Glue", bold_text)
        self.assertNotIn("SQL", bold_text)
        self.assertNotIn("patient analytics", bold_text)
        self.assertNotIn("35%", bold_text)
        self.assertNotIn("order pipelines", bold_text)
        self.assertNotIn("data quality", bold_text)

    def test_document_uses_calibri_body_text_and_bordered_title_case_headings(self) -> None:
        document = _build_sample_document()

        self.assertEqual(document.styles["Normal"].font.name, "Calibri")
        self.assertEqual(document.styles["Normal"].font.size, Pt(10))
        self.assertEqual(document.styles["List Bullet"].font.name, "Calibri")
        self.assertEqual(document.styles["List Bullet"].font.size, Pt(10))

        headings = [
            _paragraph_containing(document, "Professional Summary"),
            _paragraph_containing(document, "Technical Skills"),
            _paragraph_containing(document, "Professional Experience"),
        ]
        for heading in headings:
            self.assertEqual(heading.text, heading.text.title())
            self.assertTrue(heading.runs[0].bold)
            self.assertEqual(heading.runs[0].font.name, "Calibri")
            self.assertEqual(heading.runs[0].font.size, Pt(12))
            self.assertEqual(heading.paragraph_format.space_before, Pt(6))
            self.assertEqual(heading.paragraph_format.space_after, Pt(2))
            self.assertTrue(_has_bottom_border(heading))

    def test_experience_metadata_uses_one_compact_client_header(self) -> None:
        document = _build_sample_document()

        client_header = _paragraph_containing(
            document,
            "Example Client | Senior Data Engineer | 2023 - Present | Healthcare",
        )

        self.assertEqual(client_header.style.name, "Normal")
        self.assertEqual(client_header.paragraph_format.space_after, Pt(0))
        self.assertEqual(client_header.paragraph_format.line_spacing, 1.0)
        self.assertEqual(client_header.runs[0].text, "Example Client")
        self.assertTrue(client_header.runs[0].bold)
        self.assertEqual(client_header.runs[0].font.size, Pt(10))
        self.assertTrue(all(run.font.size == Pt(10) for run in client_header.runs[1:] if run.text))

    def test_personal_details_render_at_top_and_education_heading_is_14pt(self) -> None:
        document = _build_sample_document()

        self.assertEqual(document.paragraphs[0].text, "Nikhil Uppari")
        self.assertTrue(document.paragraphs[0].runs[0].bold)
        self.assertEqual(document.paragraphs[0].runs[0].font.size, Pt(10))
        self.assertEqual(document.paragraphs[1].text, "nikhil@example.com | 555-123-4567 | LinkedIn")
        self.assertEqual(document.paragraphs[1].runs[0].font.size, Pt(10))

        education = _paragraph_containing(document, "Education")
        self.assertEqual(education.runs[0].font.size, Pt(14))
        education_detail = _paragraph_containing(document, "M.S. Data Engineering")
        self.assertEqual(education_detail.runs[0].font.size, Pt(10))

    def test_bullets_use_tight_single_spacing(self) -> None:
        document = _build_sample_document()

        bullet = _paragraph_containing(document, "patient analytics with 35% faster order pipelines")

        self.assertEqual(bullet.style.name, "List Bullet")
        self.assertEqual(bullet.paragraph_format.space_after, Pt(0))
        self.assertEqual(bullet.paragraph_format.line_spacing, 1.0)
        self.assertAlmostEqual(bullet.paragraph_format.left_indent, Inches(0.22), delta=300)
        self.assertAlmostEqual(bullet.paragraph_format.first_line_indent, Inches(-0.16), delta=300)


def _build_sample_document():
    resume = TailoredResume(
        summary=["Built governed data platforms."],
        personal_details=["Nikhil Uppari", "nikhil@example.com | 555-123-4567 | LinkedIn"],
        technical_skills={
            "Cloud Platforms": ["AWS", "Glue", "patient analytics", "35%", "order pipelines"],
            "Data Engineering & ETL": ["ETL", "batch pipelines", "incremental loads"],
        },
        experiences=[
            Experience(
                client_name="Example Client",
                title="Senior Data Engineer",
                dates="2023 - Present",
                domain="Healthcare",
                responsibilities=[
                    "Built AWS Glue pipelines for patient analytics with 35% faster order pipelines and SQL checks.",
                ],
                environment=["AWS", "Glue", "patient analytics", "35%", "order pipelines", "SQL", "data quality"],
            )
        ],
        certifications=[],
        education="M.S. Data Engineering",
        ats_score={},
    )
    return Document(BytesIO(build_docx(resume)))


def _section_paragraphs(document, start: str, end: str):
    paragraphs = []
    in_section = False
    for paragraph in document.paragraphs:
        if paragraph.text == start:
            in_section = True
            continue
        if paragraph.text == end:
            break
        if in_section and paragraph.text:
            paragraphs.append(paragraph)
    return paragraphs


def _paragraph_containing(document, text: str):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise AssertionError(f"Could not find paragraph containing {text!r}")


def _bold_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs if run.bold is True)


def _has_bottom_border(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        return False
    return p_bdr.find(qn("w:bottom")) is not None


if __name__ == "__main__":
    unittest.main()
