from __future__ import annotations

from io import BytesIO
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from utils.schema import TailoredResume
from utils.technology_terms import is_known_technology, iter_tool_matches


BODY_FONT_NAME = "Calibri"
BODY_FONT_SIZE = Pt(10)
HEADING_FONT_SIZE = Pt(12)
EDUCATION_HEADING_FONT_SIZE = Pt(14)
SUBHEADING_FONT_SIZE = Pt(10)
SECTION_HEADING_BORDER_COLOR = "808080"
HEADING_SPACE_BEFORE = Pt(6)
HEADING_SPACE_AFTER = Pt(2)
COMPACT_SPACE_AFTER = Pt(0)
SECTION_END_SPACE_AFTER = Pt(4)
SINGLE_LINE_SPACING = 1.0
BULLET_LEFT_INDENT = Inches(0.22)
BULLET_FIRST_LINE_INDENT = Inches(-0.16)


def build_docx(resume: TailoredResume) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    for style_name in ("Normal", "List Bullet"):
        styles[style_name].font.name = BODY_FONT_NAME
        styles[style_name].font.size = BODY_FONT_SIZE

    for index, detail in enumerate(resume.personal_details):
        paragraph = document.add_paragraph()
        _compact_paragraph(paragraph)
        _add_run(paragraph, detail, bold=(index == 0))

    _heading(document, "Professional Summary")
    for item in resume.summary:
        _bullet(document, item)

    _heading(document, "Technical Skills")
    for heading, values in resume.technical_skills.items():
        if values:
            paragraph = document.add_paragraph()
            _compact_paragraph(paragraph)
            _add_run(paragraph, f"{heading}: ", bold=True, size=SUBHEADING_FONT_SIZE)
            _add_values_run(paragraph, values, bold_tools=False)

    _heading(document, "Professional Experience")
    for exp in resume.experiences:
        paragraph = document.add_paragraph()
        _compact_paragraph(paragraph)
        _add_run(paragraph, exp.client_name, bold=True, size=SUBHEADING_FONT_SIZE)
        for value in [exp.title, exp.dates, exp.domain]:
            if value:
                _add_run(paragraph, f" | {value}")

        for bullet in exp.responsibilities:
            _bullet_with_bold_tools(document, bullet, exp.environment)

        if exp.environment:
            paragraph = document.add_paragraph()
            _compact_paragraph(paragraph, space_after=SECTION_END_SPACE_AFTER)
            _add_run(paragraph, "Environment: ", bold=True, size=SUBHEADING_FONT_SIZE)
            _add_values_run(paragraph, exp.environment, bold_tools=False)

    if resume.certifications:
        _heading(document, "Certifications")
        for cert in resume.certifications:
            _bullet(document, cert)

    if resume.education:
        _heading(document, "Education", size=EDUCATION_HEADING_FONT_SIZE)
        for line in resume.education.splitlines():
            if line.strip():
                paragraph = document.add_paragraph()
                _compact_paragraph(paragraph)
                _add_run(paragraph, line.strip())

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _heading(document: Document, text: str, *, size=HEADING_FONT_SIZE) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = HEADING_SPACE_BEFORE
    paragraph.paragraph_format.space_after = HEADING_SPACE_AFTER
    paragraph.paragraph_format.line_spacing = SINGLE_LINE_SPACING
    _add_heading_border(paragraph)
    _add_run(paragraph, text, bold=True, size=size)


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _compact_bullet_paragraph(paragraph)
    _add_run(paragraph, text)


def _bullet_with_bold_tools(document: Document, text: str, tools: Iterable[str]) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _compact_bullet_paragraph(paragraph)
    matches = []
    candidates = [tool for tool in set(tools) if is_known_technology(tool)]
    for tool in sorted(candidates, key=len, reverse=True):
        for match in iter_tool_matches(text, tool):
            start, end = match.span()
            if not any(start < existing_end and end > existing_start for existing_start, existing_end, _ in matches):
                matches.append((start, end, text[start:end]))
            break
    matches.sort(key=lambda item: item[0])

    cursor = 0
    for start, end, matched_text in matches:
        if cursor < start:
            _add_run(paragraph, text[cursor:start])
        _add_run(paragraph, matched_text, bold=True)
        cursor = end
    if cursor < len(text):
        _add_run(paragraph, text[cursor:])


def _add_values_run(paragraph, values: Iterable[str], *, bold_tools: bool) -> None:
    values = [value for value in values if value]
    for index, value in enumerate(values):
        if index:
            _add_run(paragraph, ", ")
        run = _add_run(paragraph, value)
        if bold_tools and is_known_technology(value):
            run.bold = True


def _add_run(paragraph, text: str, *, bold: bool = False, size=BODY_FONT_SIZE):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT_NAME
    run.font.size = size
    return run


def _compact_paragraph(paragraph, *, space_after=COMPACT_SPACE_AFTER) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.line_spacing = SINGLE_LINE_SPACING


def _compact_bullet_paragraph(paragraph) -> None:
    _compact_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = BULLET_LEFT_INDENT
    paragraph.paragraph_format.first_line_indent = BULLET_FIRST_LINE_INDENT


def _add_heading_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), SECTION_HEADING_BORDER_COLOR)
